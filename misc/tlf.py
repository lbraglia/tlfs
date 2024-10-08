import os

import numpy as np
import pandas as pd
import docx

from dataclasses import dataclass, field
from multimethod import multimethod
from itertools import chain

debug = True

def _def_quant_display():
    return ["n", "NA",
            "min", "25pct", "75pct", "max",
            "median", "iqr",
            "mean", "sd"]

@dataclass
class Quant:
    desc: str          # long variable description
    unit: str = None   # unit of measure
    # min: float = None  # minimum
    # max: float = None  # maximum
    display: list[str] = field(default_factory = _def_quant_display)
    cell_content: str = "x"  # default cell content (single number)

    def __post_init__(self):
        # normalize display so that user can specify a single string
        if type(self.display) is str:
            self.display = [self.display]
    
    def __mul__(self, other):
        return Table(self, other)

@dataclass
class Quali:
    desc: str   # variable description
    groups: list[str]
    display: str = "n (col %)"
    cell_content: str = "x (x)"  # default cell content (single number)
    add_NA: bool = True
    add_tot: bool = True

    def __post_init__(self):
        # check quali has at least two groups
        if len(self.groups) < 2:
            raise Exception("At least two groups are required")
        # add NA and Tot to the group
        add_groups = ["NA"] * self.add_NA + ["Tot"] * self.add_tot
        self.actual_groups = self.groups + add_groups
        
    def __mul__(self, other):
        return Table(self, other)

@dataclass
class Itemset:
    """Used for Listings: basically items are the rows and contents
    the columns of a (univariate) Listing.
    """
    desc: str   # variable description
    items: list[str]     # think row
    contents: list[str]  # think columns
    cell_content: str = "x"  # default cell content (single number)

    def __post_init__(self):
        # normalize items and contents so that user can specify a
        # single string
        if type(self.items) is str:
            self.items = [self.items]
        if type(self.contents) is str:
            self.contents = [self.contents]

    def __mul__(self, other):
        return Table(self, other)
    

@dataclass
class Table():
    x: Quant|Quali
    y: Quant|Quali|None = None # none if univariate
    caption: str|None = None
    
    def __post_init__(self):
        # check input
        if self.x is None:
            raise Exception("x cannot be None")
               
        # generate the pd.DataFrame representing the table
        self._make_df(self.x, self.y)
        if debug:
            self._print_info()
            print(self.df)
        
    def to_csv(self, f):
        self.df.to_csv(f, header = False, index = False)
        
    @multimethod
    def _make_df(self, a, b):
        ta = str(type(a))
        tb = str(type(b))
        msg = "Unhandled types combination of x ({0}) and y ({1})".format(ta, tb)
        raise Exception(msg)

    @_make_df.register
    def _make_df(self, a: Quant, b: None):
        # table infos
        self.tabtype = "univariate quant"
        if self.caption is None:
            self.caption = a.desc
        self.header_nrows = 1
        self.nrows = self.header_nrows + len(a.display)
        self.ncols = 2
        # table creation
        main_header = "{0} ({1})".format(a.desc, a.unit) if a.unit is not None else a.desc
        col1 = pd.Series([""] + a.display)
        col2 = pd.Series([main_header] + [a.cell_content]*len(a.display))
        self.df = pd.concat([col1, col2], axis = 1)
        
    @_make_df.register
    def _make_df(self, a: Quali, b: None):
        # table infos
        if self.caption is None:
            self.caption = a.desc
        self.tabtype = "univariate quali"
        self.header_nrows = 1
        self.nrows = self.header_nrows + len(a.actual_groups) 
        self.ncols = 2
        # table creation
        col2_head = a.desc + ", "+ a.display
        col1 = pd.Series([""] + a.actual_groups)
        col2 = pd.Series([col2_head] + [a.cell_content]*len(a.actual_groups))
        self.df = pd.concat([col1, col2], axis = 1)

    @_make_df.register
    def _make_df(self, a: Quant, b: Quali):
        # table infos
        self.tabtype = "quant x quali"
        self.header_nrows = 2
        self.nrows = self.header_nrows + len(a.display)
        self.ncols = 1 + len(b.groups) + 2
        if self.caption is None:
            self.caption = "{0} by {1}".format(a.desc, b.desc.lower())
        # table creation
        y_header = b.desc
        x_header = "{0} ({1})".format(a.desc, a.unit) if a.unit is not None else a.desc
        header_row1 = ["", y_header] + [""] * (len(b.actual_groups) - 1)
        header_row2 = [x_header] + b.actual_groups
        self.merged_cells = [([0, 0], [1, 0]),
                             ([0, 1], [0, self.ncols - 1])]
        df   = pd.DataFrame(header_row1).transpose()
        row2 = pd.DataFrame(header_row2).transpose()
        df = pd.concat([df, row2])
        # fill the table rowwise
        for x in a.display:
            content = [x] + [a.cell_content] * (self.ncols - 1)
            new_row = pd.DataFrame(content).transpose()
            df = pd.concat([df, new_row])
        self.df = df

    @_make_df.register
    def _make_df(self, a: Quali, b: Quali):
        # table infos
        self.tabtype = "quali x quali"
        self.header_nrows = 2
        self.nrows = self.header_nrows + len(a.actual_groups)
        self.ncols = 1 + len(b.actual_groups)
        if self.caption is None:
            self.caption = "{0} by {1}".format(a.desc, b.desc.lower())
        # table creation
        y_header = b.desc + ", " + b.display
        x_header = a.desc
        header_row1 = ["", y_header] + [""] * (len(b.actual_groups) - 1)
        header_row2 = [x_header] + b.actual_groups
        self.merged_cells = [([0, 0], [1, 0]),
                             ([0, 1], [0, self.ncols - 1])]
        df   = pd.DataFrame(header_row1).transpose()
        row2 = pd.DataFrame(header_row2).transpose()
        df = pd.concat([df, row2])
        # fill the table rowwise
        for x in a.actual_groups:
            content = [x] + [b.cell_content] * (self.ncols - 1)
            new_row = pd.DataFrame(content).transpose()
            df = pd.concat([df, new_row])
        self.df = df

    @_make_df.register
    def _make_df(self, a: Itemset, b: None):
        # table infos
        self.tabtype = "univariate listing"
        self.header_nrows = 1
        self.nrows = self.header_nrows + len(a.items)
        self.ncols = 1 + len(a.contents)
        if self.caption is None:
            self.caption = "Listing of {0}".format(a.desc.lower())
        # table creation
        header_row1 = [""] + a.contents
        df   = pd.DataFrame(header_row1).transpose()
        # fill the table rowwise
        for x in a.items:
            content = [x] + [a.cell_content] * (self.ncols - 1)
            new_row = pd.DataFrame(content).transpose()
            df = pd.concat([df, new_row])
        self.df = df

    @_make_df.register
    def _make_df(self, a: Itemset, b: Quali):
        # table infos
        self.tabtype = "bivariate listing"
        self.header_nrows = 2
        self.nrows = self.header_nrows + len(a.items)
        self.ncols = 1 + len(a.contents) * len(b.groups)
        if self.caption is None:
            self.caption = "Listing of {0} by {1}".format(a.desc.lower(), b.desc.lower())
        # table creation
        head = [[x] + ([""] * (len(a.contents) - 1))    for x in b.groups]
        header_row1 = [""] + list(chain.from_iterable(head)) # flatten it
        # set the cell to be merged
        # header_row1_merged_start = list(range(2, 7, 3))
        # header_row1_merged_stop =  list(range(2 + 3 - 1, 7 + 1, 3))
        header_row1_merged_start = list(range(1, self.ncols, len(a.contents)))
        header_row1_merged_stop =  list(range(1 + len(a.contents) - 1,
                                              self.ncols + 1,
                                              len(a.contents)))
        self.merged_cells = [([0, sta], [0, sto]) for sta, sto in zip(
            header_row1_merged_start,
            header_row1_merged_stop
        )]
        if debug:
            print(self.merged_cells)
        header_row2 = [""] + a.contents * len(b.groups)
        df   = pd.DataFrame(header_row1).transpose()
        row2 = pd.DataFrame(header_row2).transpose()
        df = pd.concat([df, row2])
        # fill the table rowwise
        for x in a.items:
            content = [x] + [a.cell_content] * (self.ncols - 1)
            new_row = pd.DataFrame(content).transpose()
            df = pd.concat([df, new_row])
        self.df = df
        
    def _print_info(self):
        report = """
        caption = {0} (tabtype = {1}),
        header_nrows = {2}, nrows = {3},  ncols = {4}
        """.format(
            self.caption,
            self.tabtype,
            self.header_nrows,
            self.nrows,
            self.ncols)
        print(report)
    
    def add_to_docx(self, doc: docx.Document):
        doc.add_paragraph(self.caption)
        tab = doc.add_table(rows = self.nrows, cols = self.ncols)
        # tab.style = "Table Grid" #funzionasse
        # tab.alignment = WD_TABLE_ALIGNMENT.LEFT
        # tab.autofit = False
        # tab.style = None

        # add contents
        for r in range(self.nrows):
            for c in range(self.ncols):
                content = str(self.df.iat[r, c])
                if content != "":
                    tab.cell(r, c).text = content

        # merge cells if self.merged_cells is defined
        try:
            self.merged_cells
        except AttributeError:
            pass
        else:
            for cell_range in self.merged_cells:
                start = cell_range[0]
                stop = cell_range[1]
                start_cell = tab.cell(start[0], start[1])
                stop_cell = tab.cell(stop[0], stop[1])
                start_cell.merge(stop_cell)

        # blank line at the end
        doc.add_paragraph("")


@dataclass
class Section:
    title: str = ""
    tables: list[Table] = field(default_factory = list)
    heading_lev: int = 2

    def add_tables(self, x: Table|list[Table]):
        if type(x) == list:
            self.tables.expand(x)
        elif type(x) == Table:
            self.tables.append(x)
        else:
            raise Exception("x must be a Table or a list of tables")
        
    def add_to_docx(self, doc: docx.Document):
        # heading and space
        doc.add_heading(self.title, level = self.heading_lev)
        doc.add_paragraph("")
        # add tables
        for t in self.tables:
            t.add_to_docx(doc)

            
@dataclass
class TLF:
    title: str|None = None
    sections: list[Section] = field(default_factory = list)
    heading_lev: int = 1

    def add_sections(self, x: Section|list[Section]):
        if type(x) == list:
            self.sections.expand(x)
        elif type(x) == Section:
            self.sections.append(x)
        else:
            raise Exception("x must be a Section or a list of sections")
    
    def to_docx(self, outfile: str = "/tmp/test.docx", view: bool = True):
        doc = docx.Document()
        # header
        if type(self.title) is str:
            doc.add_heading(self.title, level = self.heading_lev)
            
        # content
        for s in self.sections:
            s.add_to_docx(doc)

        doc.save(outfile)
        if view:
            os.system("libreoffice " + outfile)

    def from_xlsx(self, infile = "tlf_structure.xlsx"):
        sections = pd.read_excel(infile, sheet_name = 'sections')
        tables = pd.read_excel(infile, sheet_name = 'tables')
        variables = pd.read_excel(infile, sheet_name = 'variables')
        groups_items_contents = pd.read_excel(infile, sheet_name = 'groups_items_contents')

        if debug:
            print(sections)
            print(tables)
            print(variables)
            print(groups_items_contents)

        # already created variables
        variables_pool = {} 

        # groups_items_contents as lookup dict
        gic = {}
        for _, gic_id, gic_gic in groups_items_contents.itertuples():
            if gic_id not in gic:
                # create a new gic
                gic[gic_id] = [gic_gic]
            else:
                # update
                gic[gic_id].append(gic_gic)

        if debug:
            print(gic)

        # function to create variables
        def make_var(varid):
            vrow = variables.loc[variables.id == varid]
            if vrow.shape[0] > 1:
                msg = "multiple variables with id {0}".format(varid)
                raise Exception(msg)

            for v in vrow.itertuples(): # avoid squeezing all the var with a 1 iteration for
                if v.type == "quant":
                    unit = None if v.unit == "" else v.unit
                    return Quant(desc = v.desc, unit = unit)
                elif v.type == "quali":
                    groups = gic[v.groups]
                    return Quali(desc = v.desc, groups = groups)
                elif v.type == "itemset":
                    items = gic[v.items]
                    contents = gic[v.contents]
                    return Itemset(desc = v.desc, items = items, contents = contents)
                else:
                    msg = "Unhandled variable type: '{0}'".format(v.type) 
                    raise Exception(msg)

        
        # for every section
        for _, sect_id, sect_title in sections.itertuples():

            # initialize the data structure
            sect = Section(sect_title)
            # Select tables for this section
            sect_tables = tables.loc[tables.section == sect_id, ] 

            # # for every table in the section construct the
            for _, _, tab_var1, tab_var2, tab_caption in sect_tables.itertuples():

                # retrieve or create the variables and put them in their container
                # tab_var1 is mandatory
                if tab_var1 in (np.nan, ""):
                    raise Exception("var1 cannot be missing")
                elif tab_var1 not in variables_pool: # still not encountered variables
                    used_var1 = variables_pool[tab_var1] = make_var(tab_var1)
                else: # already encountered
                    used_var1 = variables_pool[tab_var1]
                
                # tab_var2 is optional
                if tab_var2 in (np.nan, ""):
                    used_var2 = None
                elif tab_var2 not in variables_pool:
                    used_var2 = variables_pool[tab_var2] = make_var(tab_var2)
                else:
                    used_var2 = variables_pool[tab_var2]

                # normalize caption
                if tab_caption in ("", np.nan):
                    used_caption = None
                else:
                    used_caption = tab_caption
                
                # add the table to the section
                sect.add_tables(Table(x = used_var1,
                                      y = used_var2,
                                      caption = used_caption))

            # add the section to the tlf
            self.add_sections(sect)
 
               

tlf = TLF("Test")
tlf.from_xlsx()
tlf
tlf.to_docx(view = True)


    
# for custom shit do program by hand
# ----------------------------------

# overloading can be useful: * in variables to create Tables
age = Quant("Age")
sex = Quali("Sex", groups = ["M", "F"])
trt = Quali("Treatment", groups = ["EXP", "CTRL"])
prices = Itemset("Unit costs",
                 items = ["Dentist", "Hospice", "Blood test"],
                 contents = ["unit cost", "per", "source"])
nation = Quali("Nation", groups = ["UK", "ITA"])

age * trt
[var * trt for var in [age, sex]] + [prices * nation]


# changing default example
age2 = Quant("Age", display = ["median", "25pct", "75pct"], unit = 'years')
age3 = Quant("Age", display = "median (iqr)", cell_content = "xx (xx - xx)")
sex2 = Quali("Sex", groups = ["M", "F"], display = "n", cell_content = "x")


univ = Section("Univariate tables", [Table(age), Table(sex)])
changed_def = Section("Some changed defaults", [Table(age2), age2 * trt, age3 * trt, Table(sex2)])
listings = Section("Listings", [Table(prices), prices * nation])
tlf2 = TLF("Table, Listings, Figure examples", [univ, biv, changed_def, listings])
tlf.to_doc()


# todo instead of
univ = Section("Univariate tables", [Table(age), Table(sex)])
# just do
univ = Section("Univariate tables", [age, sex])


## aggiungere in post init e nei method che aggiungono cose alle sezioni direi
